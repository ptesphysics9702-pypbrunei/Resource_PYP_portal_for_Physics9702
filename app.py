import io
import os
import re
import fitz  # PyMuPDF
import streamlit as st

# Word Document Libraries
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Google API Libraries using Service Account
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from googleapiclient.errors import HttpError

# ==========================================
# 1. CONFIGURATION & DRIVE FOLDER MAPPING (PHYSICS 9702)
# ==========================================
SYLLABUS_CODE = "9702"

# Google Drive Folder IDs mapped to live Google Drive folders
FOLDER_IDS = {
    "theory": "180eJGPSt53c0u3Fx-39G8ltcopB2eT-b",     # Papers 1, 2, 4, 5
    "practical": "10yymDTMshSmyjBB5VGhRp6TjGbS3QsK7",  # Paper 3 (33, 34, 35, 36)
    "zips": "17I6Cm9H_BMfP_TMoqurTjay-29BkD1km",       # Confidential Instructions (_ci_) / Data ZIPs (_sf_)
    "ms": "1McFgLCC5xB7G6KOwUnsIwN5bKpfGVdMD"          # Dedicated Marking Schemes Folder
}

# Local directories for mirroring files on the server
LOCAL_FOLDERS = {
    "theory": f"{SYLLABUS_CODE}_theory",
    "practical": f"{SYLLABUS_CODE}_practical",
    "zips": f"{SYLLABUS_CODE}_zips",
    "ms": f"{SYLLABUS_CODE}_ms"
}

# Ensure local storage directories exist on server startup
for folder_path in LOCAL_FOLDERS.values():
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)

# ==========================================
# 2. STATE INITIALIZATION & CALLBACKS
# ==========================================
if 'handout_basket' not in st.session_state:
    st.session_state.handout_basket = []
if 'theory_results' not in st.session_state:
    st.session_state.theory_results = []
if 'practical_results' not in st.session_state:
    st.session_state.practical_results = []
if 'sync_status' not in st.session_state:
    st.session_state.sync_status = None

def add_to_basket(item):
    """Callback to safely add an item to the basket before UI renders."""
    st.session_state.handout_basket.append(item)
    st.toast("Added to basket!", icon="🛒")

def remove_from_basket(index):
    """Callback to safely remove an item from the basket before UI renders."""
    if 0 <= index < len(st.session_state.handout_basket):
        st.session_state.handout_basket.pop(index)
        st.toast("Item removed from basket.", icon="🗑️")

def clear_basket():
    """Callback to clear all items from the basket."""
    st.session_state.handout_basket = []
    st.toast("Basket cleared!", icon="🧹")

# ==========================================
# 3. HELPER FUNCTIONS: WORD DOCUMENT XML
# ==========================================
def add_page_number_to_run(run):
    """
    Inserts a dynamic Word PAGE field into a document text run using OpenXML.
    """
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

# ==========================================
# 4. SERVICE ACCOUNT GOOGLE DRIVE API ENGINE
# ==========================================
SCOPES = ["https://www.googleapis.com/auth/drive"]

@st.cache_resource(ttl=3500)
def get_shared_drive_service():
    """
    Authenticates with Google Drive API using a Google Service Account stored in Streamlit Secrets.
    """
    try:
        if "gcp_service_account" in st.secrets:
            creds_info = dict(st.secrets["gcp_service_account"])
        else:
            creds_info = dict(st.secrets)

        creds = service_account.Credentials.from_service_account_info(
            creds_info, scopes=SCOPES
        )
        return build('drive', 'v3', credentials=creds)

    except KeyError as ke:
        st.error(f"❌ Missing Google Service Account Secret Key: {ke}")
        return None
    except Exception as e:
        st.error(f"❌ Service Account Authentication Error: {e}")
        return None

def build_drive_service():
    """Retrieves cached Google Drive client instance."""
    return get_shared_drive_service()

def sync_drive_folder_to_local(folder_key: str) -> tuple[int, str]:
    """
    Queries Google Drive for a folder and downloads missing files to local server mirror.
    """
    service = build_drive_service()
    if not service:
        return 0, "Failed to authenticate with Google Drive Service Account."

    drive_folder_id = FOLDER_IDS[folder_key]
    local_path = LOCAL_FOLDERS[folder_key]

    try:
        query = f"'{drive_folder_id}' in parents and trashed = false"
        results = service.files().list(q=query, fields="files(id, name, mimeType)").execute()
        drive_files = results.get('files', [])

        downloaded_count = 0

        for file_info in drive_files:
            file_name = file_info['name']
            file_id = file_info['id']
            local_file_path = os.path.join(local_path, file_name)

            if not os.path.exists(local_file_path):
                request = service.files().get_media(fileId=file_id)
                with open(local_file_path, "wb") as f:
                    downloader = MediaIoBaseDownload(f, request)
                    done = False
                    while not done:
                        status, done = downloader.next_chunk()
                downloaded_count += 1

        return downloaded_count, f"Synced {downloaded_count} new file(s)"

    except HttpError as http_err:
        return 0, f"Sync HTTP Error ({http_err.resp.status})"
    except Exception as e:
        return 0, f"Sync error: {e}"

# ==========================================
# 5. CACHED AUTOMATIC DRIVE SYNC ENGINE
# ==========================================
@st.cache_data(ttl=3600, show_spinner="🔄 Auto-syncing Google Drive files...")
def auto_sync_all_folders():
    """
    Automatically checks and syncs all 4 Google Drive folders (theory, practical, zips, ms).
    """
    total_new = 0
    sync_details = {}
    for f_key in ["theory", "practical", "zips", "ms"]:
        count, msg = sync_drive_folder_to_local(f_key)
        total_new += count
        sync_details[f_key] = msg
    return total_new, sync_details

# ==========================================
# 6. SEARCH ENGINE & HELPER FUNCTIONS
# ==========================================
def render_pdf_page_preview(filepath: str, page_num: int):
    """Renders a single PDF page into PNG bytes for Streamlit previewing."""
    try:
        doc = fitz.open(filepath)
        page = doc.load_page(page_num)
        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
        img_bytes = pix.tobytes("png")
        doc.close()
        return img_bytes
    except Exception as e:
        st.error(f"Unable to render page preview: {e}")
        return None

def search_pdfs(keyword_list, folder_path, allowed_variants, match_mode="ALL"):
    """Scans local PDF files for keywords using flexible matching."""
    results = []
    if not os.path.exists(folder_path):
        return results

    cleaned_keywords = [k.strip().lower() for k in keyword_list if k.strip()]
    if not cleaned_keywords:
        return results

    for file in os.listdir(folder_path):
        if file.endswith(".pdf"):
            base_name = os.path.splitext(file)[0]
            
            if "_ci_" in file:
                continue

            is_valid_variant = any(base_name.endswith(f"_{variant}") for variant in allowed_variants)
            if not is_valid_variant:
                continue

            filepath = os.path.join(folder_path, file)
            try:
                doc = fitz.open(filepath)
                for page_num in range(len(doc)):
                    page_text = doc[page_num].get_text()
                    matched_keywords_count = 0
                    
                    for kw in cleaned_keywords:
                        escaped_kw = re.escape(kw)
                        pattern = r'\b' + escaped_kw + r'(s|es)?\b'
                        
                        if re.search(pattern, page_text, re.IGNORECASE) or kw in page_text.lower():
                            matched_keywords_count += 1

                    if match_mode == "ALL" and matched_keywords_count == len(cleaned_keywords):
                        results.append({
                            "file": file,
                            "page": page_num,
                            "path": filepath,
                            "type": "QP" if "_qp_" in file else "MS"
                        })
                    elif match_mode == "ANY" and matched_keywords_count > 0:
                        results.append({
                            "file": file,
                            "page": page_num,
                            "path": filepath,
                            "type": "QP" if "_qp_" in file else "MS"
                        })
                        
                doc.close()
            except Exception:
                continue
                
    return results

# ==========================================
# 7. STREAMLIT UI LAYOUT & STYLING
# ==========================================
st.set_page_config(page_title="9702 Physics PYP Archives", layout="wide")

MAIN_BG_COLOR = "#FEE7F9"     # Soft Pink
SIDEBAR_BG_COLOR = "#FCBBEF"  # Accent Pink
INPUT_BAR_COLOR = "#E7FCBB"   # Soft Green

st.markdown(
    f"""
    <style>
    .stAppViewContainer {{ background-color: {MAIN_BG_COLOR}; }}
    .stHeader {{ background-color: {MAIN_BG_COLOR}; }}
    [data-testid="stSidebar"] {{ background-color: {SIDEBAR_BG_COLOR}; }}
    div[data-testid="stTextInput"] input, div[data-baseweb="input"] {{
        background-color: {INPUT_BAR_COLOR} !important; border-radius: 8px !important;
    }}
    div[data-testid="stSelectbox"] div {{
        background-color: {INPUT_BAR_COLOR} !important; border-radius: 8px !important;
    }}
    </style>
    """,
    unsafe_allow_html=True
)

st.title("GCE A/AS LEVEL PHYSICS")
st.subheader("⚡ 9702 Physics PYP Resource Library")

# ==========================================
# 8. SIDEBAR CONTROL PANEL
# ==========================================
with st.sidebar:
    st.header("⚡ Sync Control Panel")
    
    # Manual Sync Button with Immediate On-Screen Indicator
    if st.button("🔄 Sync Drive Files Now", use_container_width=True):
        with st.spinner("Connecting to Google Drive..."):
            st.cache_data.clear()
            total_new, sync_details = auto_sync_all_folders()
            st.session_state.sync_status = {
                "total": total_new,
                "details": sync_details
            }

    # Render Sync Indicator & Status Breakdown Directly Below the Button
    if st.session_state.sync_status is not None:
        total_synced = st.session_state.sync_status["total"]
        details = st.session_state.sync_status["details"]

        if total_synced > 0:
            st.success(f"✅ **{total_synced}** new file(s) downloaded!")
        else:
            st.info("ℹ️ Sync complete. Local files are up to date.")

        st.markdown("**Sync Details per Folder:**")
        st.caption(f"• **Theory:** {details.get('theory', 'Up to date')}")
        st.caption(f"• **Practical:** {details.get('practical', 'Up to date')}")
        st.caption(f"• **Marking Schemes:** {details.get('ms', 'Up to date')}")
        st.caption(f"• **Instructions:** {details.get('zips', 'Up to date')}")

    st.markdown("---")
    
    # Handout Basket Summary
    st.header("Handout Basket Summary")
    st.metric(label="Saved Pages in Basket", value=len(st.session_state.handout_basket))
    st.button("🗑️ Clear Entire Basket", use_container_width=True, on_click=clear_basket)

# ==========================================
# 9. NAVIGATION TABS
# ==========================================
tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "🔍 Theory Search", 
    "⚙️ Practical Search", 
    "🛒 Handout Cart", 
    "🔑 Mark Scheme",
    "📦 PYP Instruction", 
    "🔒 Upload PYP"
])

# --- TAB 1: THEORY SEARCH ---
with tab1:
    st.header("Search Physics Theory Papers")
    st.caption("Variants: P1 (12, 13) | P2 (22, 23) | P4 (42, 43) | P5 (52, 53)")
    
    col_t1_kw, col_t1_mode = st.columns([3, 1])
    with col_t1_kw:
        keyword_t1 = st.text_input("Enter Theory Keywords (comma-separated)", placeholder="e.g., resistor, velocity, quantum", key="t1_kw")
    with col_t1_mode:
        match_mode_t1 = st.selectbox("Search Match Mode", options=["Match ALL (AND)", "Match ANY (OR)"], key="t1_mode")

    if st.button("Search Theory Papers", type="primary"):
        if keyword_t1.strip():
            with st.spinner("Scanning Theory PDFs..."):
                keywords = [k.strip() for k in keyword_t1.split(",") if k.strip()]
                theory_variants = ["12", "13", "22", "23", "42", "43", "52", "53"]
                selected_mode = "ALL" if "ALL" in match_mode_t1 else "ANY"
                st.session_state.theory_results = search_pdfs(keywords, LOCAL_FOLDERS["theory"], theory_variants, match_mode=selected_mode)
        else:
            st.warning("Please enter at least one keyword.")

    if st.session_state.theory_results:
        st.write(f"Found **{len(st.session_state.theory_results)}** matching pages:")
        for idx, item in enumerate(st.session_state.theory_results):
            doc_kind = "📝 Question Paper" if item["type"] == "QP" else "🔑 Marking Scheme"
            with st.expander(f"📄 {item['file']} | {doc_kind} | Page {item['page'] + 1}"):
                c1, c2 = st.columns([3, 1])
                with c1:
                    preview_img = render_pdf_page_preview(item["path"], item["page"])
                    if preview_img:
                        st.image(preview_img, caption=f"Preview Page {item['page'] + 1}", use_container_width=True)
                with c2:
                    st.button("➕ Add to Basket", key=f"add_t1_{idx}", on_click=add_to_basket, args=(item,))
                    if os.path.exists(item["path"]):
                        with open(item["path"], "rb") as pdf_file:
                            st.download_button("📥 Download Full PDF", data=pdf_file, file_name=item["file"], mime="application/pdf", key=f"dl_t1_{idx}")


# --- TAB 2: PRACTICAL SEARCH ---
with tab2:
    st.header("Search Physics Practical Papers")
    st.caption("Variants: Paper 3 (33, 34, 35, 36)")
    
    col_t2_kw, col_t2_mode = st.columns([3, 1])
    with col_t2_kw:
        keyword_t2 = st.text_input("Enter Practical Keywords (comma-separated)", placeholder="e.g., oscillation, resistance, uncertainty", key="t2_kw")
    with col_t2_mode:
        match_mode_t2 = st.selectbox("Search Match Mode", options=["Match ALL Keywords", "Match ANY Keyword"], key="t2_mode")

    if st.button("Search Practical Papers", type="primary"):
        if keyword_t2.strip():
            with st.spinner("Scanning Practical PDFs..."):
                keywords = [k.strip() for k in keyword_t2.split(",") if k.strip()]
                practical_variants = ["33", "34", "35", "36"]
                selected_mode = "ALL" if "ALL" in match_mode_t2 else "ANY"
                st.session_state.practical_results = search_pdfs(keywords, LOCAL_FOLDERS["practical"], practical_variants, match_mode=selected_mode)
        else:
            st.warning("Please enter at least one keyword.")

    if st.session_state.practical_results:
        st.write(f"Found **{len(st.session_state.practical_results)}** matching pages:")
        for idx, item in enumerate(st.session_state.practical_results):
            doc_kind = "📝 Question Paper" if item["type"] == "QP" else "🔑 Marking Scheme"
            with st.expander(f"📄 {item['file']} | {doc_kind} | Page {item['page'] + 1}"):
                c1, c2 = st.columns([3, 1])
                with c1:
                    preview_img = render_pdf_page_preview(item["path"], item["page"])
                    if preview_img:
                        st.image(preview_img, caption=f"Preview Page {item['page'] + 1}", use_container_width=True)
                with c2:
                    st.button("➕ Add to Basket", key=f"add_t2_{idx}", on_click=add_to_basket, args=(item,))
                    if os.path.exists(item["path"]):
                        with open(item["path"], "rb") as pdf_file:
                            st.download_button("📥 Download Full PDF", data=pdf_file, file_name=item["file"], mime="application/pdf", key=f"dl_t2_{idx}")


# --- TAB 3: HANDOUT CART & WORD BUILDER ---
with tab3:
    st.header("Worksheet / Handout Builder")
    if st.session_state.handout_basket:
        st.subheader(f"Selected Pages: {len(st.session_state.handout_basket)}")

        for idx, item in enumerate(st.session_state.handout_basket):
            col_info, col_remove = st.columns([4, 1])
            col_info.write(f"{idx + 1}. **{item['file']}** (Page {item['page'] + 1})")
            col_remove.button("❌ Remove", key=f"remove_basket_{idx}", on_click=remove_from_basket, args=(idx,))

        st.markdown("---")
        if st.button("🪄 Export Handout to Word Document", type="primary"):
            try:
                doc = Document()
                section = doc.sections[0]

                section.orientation = WD_ORIENT.PORTRAIT
                section.page_width = Inches(8.5)
                section.page_height = Inches(11.5)
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.3)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)

                header = section.header
                header_paragraph = header.paragraphs[0]
                header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                header_run = header_paragraph.add_run("Page ")
                header_run.font.name = "Calibri"
                header_run.font.size = Pt(10)
                add_page_number_to_run(header_run)

                doc.add_heading(f'PTES {SYLLABUS_CODE} Physics Handout', level=1)

                for i, item in enumerate(st.session_state.handout_basket):
                    doc.add_heading(f"Source: {item['file']} (Page {item['page'] + 1})", level=2)
                    pdf_doc = fitz.open(item['path'])
                    page = pdf_doc.load_page(item['page'])
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img_data = io.BytesIO(pix.tobytes("png"))
                    
                    doc.add_picture(img_data, width=Inches(6.8))
                    
                    if i < len(st.session_state.handout_basket) - 1:
                        doc.add_page_break()
                        
                    pdf_doc.close()

                target_filename = f"{SYLLABUS_CODE}_Physics_Handout.docx"
                doc.save(target_filename)

                with open(target_filename, "rb") as f:
                    st.download_button("📥 Click to Download Word Handout", data=f, file_name=target_filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except Exception as e:
                st.error(f"❌ Handout Export Failed: {e}")
    else:
        st.info("Your basket is empty. Add pages from Tab 1 or Tab 2.")


# --- TAB 4: MARKING SCHEME LOOKUP TAB ---
with tab4:
    st.header("🔑 Download Marking Schemes (MS)")
    st.caption("Select Year, Session, and Variant to quickly preview and download official Marking Schemes.")
    
    col_ms1, col_ms2, col_ms3 = st.columns(3)
    
    with col_ms1:
        ms_year = st.selectbox("Select Year", [str(y) for y in range(2026, 2018, -1)], key="ms_year_select")
    with col_ms2:
        ms_session = st.selectbox("Select Session", ["May/June (s)", "October/November (w)", "March (m)"], key="ms_session_select")
        ms_session_code = ms_session.split("(")[1].replace(")", "")
    with col_ms3:
        ms_variant = st.selectbox(
            "Select Variant / Component", 
            ["12", "13", "22", "23", "33", "34", "35", "36", "42", "43", "52", "53"], 
            key="ms_variant_select"
        )

    short_ms_year = ms_year[-2:]
    target_ms_filename = f"{SYLLABUS_CODE}_{ms_session_code}{short_ms_year}_ms_{ms_variant}.pdf"

    # Search across Marking Schemes folder as well as Theory & Practical folders
    found_ms_path = None
    for folder_key in ["ms", "theory", "practical"]:
        check_path = os.path.join(LOCAL_FOLDERS[folder_key], target_ms_filename)
        if os.path.exists(check_path):
            found_ms_path = check_path
            break

    st.markdown("---")
    if found_ms_path:
        st.success(f"Found Marking Scheme: `{target_ms_filename}`")
        
        c_ms_preview, c_ms_download = st.columns([3, 1])
        
        with c_ms_preview:
            ms_preview_img = render_pdf_page_preview(found_ms_path, 0)
            if ms_preview_img:
                st.image(ms_preview_img, caption=f"Preview of {target_ms_filename} (Page 1)", use_container_width=True)
                
        with c_ms_download:
            with open(found_ms_path, "rb") as ms_f:
                st.download_button(
                    label=f"📥 Download `{target_ms_filename}`",
                    data=ms_f,
                    file_name=target_ms_filename,
                    mime="application/pdf",
                    type="primary"
                )
    else:
        st.warning(
            f"No Marking Scheme found for `{SYLLABUS_CODE}_{ms_session_code}{short_ms_year}` Component `{ms_variant}` (`{target_ms_filename}`). "
            "Click 'Sync Drive Files Now' in the sidebar if you recently uploaded it to Google Drive."
        )


# --- TAB 5: CONFIDENTIAL INSTRUCTIONS (_ci_) & SOURCE FILES ---
with tab5:
    st.header("Download Practical Confidential Instructions")
    c1, c2, c3 = st.columns(3)
    with c1:
        z_year = st.selectbox("Select Year", [str(y) for y in range(2026, 2018, -1)])
    with c2:
        z_session = st.selectbox("Select Session", ["May/June (s)", "October/November (w)"])
        session_code = z_session.split("(")[1].replace(")", "")
    with c3:
        z_paper = st.selectbox("Select Paper Component", ["Paper 33", "Paper 34", "Paper 35", "Paper 36"])

    short_year = z_year[-2:]
    possible_filenames = [
        f"{SYLLABUS_CODE}_{session_code}{short_year}_ci_{z_paper}.pdf",
        f"{SYLLABUS_CODE}_{session_code}{short_year}_ci_{z_paper}.zip",
        f"{SYLLABUS_CODE}_{session_code}{short_year}_sf_{z_paper}.zip",
        f"{SYLLABUS_CODE}_{session_code}{short_year}_sf_{z_paper}.pdf",
    ]

    found_file_path = None
    matched_filename = None

    for fname in possible_filenames:
        check_path = os.path.join(LOCAL_FOLDERS["zips"], fname)
        if os.path.exists(check_path):
            found_file_path = check_path
            matched_filename = fname
            break

    st.markdown("---")
    if found_file_path and matched_filename:
        st.success(f"Found File: `{matched_filename}`")
        mime_type = "application/pdf" if matched_filename.endswith(".pdf") else "application/zip"
        
        with open(found_file_path, "rb") as f:
            st.download_button(
                label=f"📥 Download {matched_filename}",
                data=f,
                file_name=matched_filename,
                mime="application/pdf" if mime_type == "application/pdf" else "application/zip"
            )
            
        if matched_filename.endswith(".pdf"):
            preview = render_pdf_page_preview(found_file_path, 0)
            if preview:
                st.image(preview, caption=f"Preview of {matched_filename} (Page 1)", width=600)
    else:
        st.warning(
            f"No Instructions Source File found for `{SYLLABUS_CODE}_{session_code}{short_year}` Paper `{z_paper}`. "
            "Use the Sidebar Sync button to pull newly added files."
        )


# --- TAB 6: ADMIN & DIRECT GOOGLE DRIVE ACCESS ---
with tab6:
    st.header("📂 Open Target Google Drive Folders Directly")
    st.caption("Click any button below to open the corresponding Google Drive folder in a new tab to add or manage past papers.")

    admin_password = st.secrets.get("ADMIN_PASSWORD")

    if not admin_password:
        st.error("🚨 `ADMIN_PASSWORD` is not configured in Streamlit Secrets.")
    else:
        pwd = st.text_input("Enter Admin Password to Reveal Folder Links", type="password")

        if pwd == admin_password:
            st.success("Admin Access Granted")
            
            c_btn1, c_btn2 = st.columns(2)
            
            with c_btn1:
                st.subheader("⚙️ Practical Papers Folder")
                st.caption("For Paper 3 variants (33, 34, 35, 36)")
                st.link_button(
                    "📂 Open Practical Drive Folder", 
                    f"https://drive.google.com/drive/folders/{FOLDER_IDS['practical']}", 
                    type="primary", 
                    use_container_width=True
                )
                
                st.markdown("---")
                
                st.subheader("🔑 Marking Schemes Folder")
                st.caption("For all Marking Schemes (QP variants 12–53)")
                st.link_button(
                    "📂 Open Marking Schemes Drive Folder", 
                    f"https://drive.google.com/drive/folders/{FOLDER_IDS['ms']}", 
                    type="primary", 
                    use_container_width=True
                )

            with c_btn2:
                st.subheader("🔍 Theory Papers Folder")
                st.caption("For Papers 1, 2, 4, 5 variants (12, 13, 22, 23, 42, 43, 52, 53)")
                st.link_button(
                    "📂 Open Theory Drive Folder", 
                    f"https://drive.google.com/drive/folders/{FOLDER_IDS['theory']}", 
                    type="primary", 
                    use_container_width=True
                )
                
                st.markdown("---")
                
                st.subheader("📦 Practical Instructions Folder")
                st.caption("For Confidential Instructions (_ci_) and Source Files (_sf_)")
                st.link_button(
                    "📂 Open Instructions Drive Folder", 
                    f"https://drive.google.com/drive/folders/{FOLDER_IDS['zips']}", 
                    type="primary", 
                    use_container_width=True
                )

# ==========================================
# 10. FOOTER
# ==========================================
st.markdown("---")
st.markdown(
    """
    <div style="text-align: center; width: 100%;">
        <p style="font-size: 20px; font-weight: bold; margin-bottom: 5px;">✨ Digital 9702 Physics Resource Portal ✨</p>
        <p style="color: gray; font-size: 14px;">Creator: HNHaziqah Computer Science PTES</p>
    </div>
    """,
    unsafe_allow_html=True
)
